"""
Tests for document parsing.
"""
import pytest
from app.services.parser import DocumentParser
from app.services.document_block import NoExtractableTextError


@pytest.fixture
def parser():
    return DocumentParser()


def test_get_file_type(parser):
    """Test file type detection."""
    assert parser.get_file_type("test.pdf") == "pdf"
    assert parser.get_file_type("test.docx") == "word"
    assert parser.get_file_type("test.xlsx") == "excel"
    assert parser.get_file_type("test.txt") == "text"
    assert parser.get_file_type("test.unknown") == "unknown"


def test_parse_text(parser, tmp_path):
    """Test parsing a text file."""
    text_file = tmp_path / "test.txt"
    text_file.write_text("Hello, this is a test document.\nIt has multiple lines.")

    docs, meta = parser.parse(str(text_file))

    assert len(docs) == 1
    assert "Hello, this is a test document" in docs[0].page_content
    assert meta["file_type"] == "text"
    assert meta["filename"] == "test.txt"
    assert docs[0].block_type == "prose"
    assert docs[0].metadata["document_title"] == "test"


def test_empty_text_is_rejected(parser, tmp_path):
    text_file = tmp_path / "empty.txt"
    text_file.write_text("   \n")

    with pytest.raises(NoExtractableTextError, match="NO_EXTRACTABLE_TEXT"):
        parser.parse(str(text_file))


def test_blank_pdf_is_rejected_instead_of_succeeding_with_zero_chunks(parser, tmp_path):
    import fitz

    pdf_path = tmp_path / "scan.pdf"
    pdf = fitz.open()
    pdf.new_page()
    pdf.save(pdf_path)
    pdf.close()

    with pytest.raises(NoExtractableTextError, match="NO_EXTRACTABLE_TEXT"):
        parser.parse(str(pdf_path))


def test_word_preserves_heading_and_table_order(parser, tmp_path):
    from docx import Document

    path = tmp_path / "handbook.docx"
    document = Document()
    document.add_heading("安全规范", level=1)
    document.add_paragraph("进入车间前需要完成安全培训。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "岗位"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = "操作员"
    table.cell(1, 1).text = "持证上岗"
    document.add_paragraph("培训记录由管理员归档。")
    document.save(path)

    blocks, meta = parser.parse(str(path))

    assert [block.block_type for block in blocks] == ["prose", "table", "prose"]
    assert all(block.heading_path == ("安全规范",) for block in blocks)
    assert "| 岗位 | 要求 |" in blocks[1].page_content
    assert "| 操作员 | 持证上岗 |" in blocks[1].page_content
    assert meta["block_count"] == 3


def test_word_heading_only_document_keeps_extractable_text(parser, tmp_path):
    from docx import Document

    path = tmp_path / "headings.docx"
    document = Document()
    document.add_heading("一级标题", level=1)
    document.add_heading("二级标题", level=2)
    document.save(path)

    blocks, _ = parser.parse(str(path))

    assert [block.page_content for block in blocks] == ["一级标题", "二级标题"]
    assert blocks[1].heading_path == ("一级标题", "二级标题")


def test_markdown_same_level_heading_replaces_instead_of_nesting(parser, tmp_path):
    path = tmp_path / "headings.md"
    path.write_text("## 章节 A\n内容 A\n## 章节 B\n内容 B", encoding="utf-8")

    blocks, _ = parser.parse(str(path))

    assert [block.heading_path for block in blocks] == [("章节 A",), ("章节 B",)]


def test_excel_is_a_structured_table_with_sheet_path(parser, tmp_path):
    from openpyxl import Workbook

    path = tmp_path / "inventory.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "设备清单"
    sheet.append(["设备", "状态"])
    sheet.append(["服务器", "运行中"])
    workbook.save(path)

    blocks, _ = parser.parse(str(path))

    assert len(blocks) == 1
    assert blocks[0].block_type == "table"
    assert blocks[0].heading_path == ("设备清单",)
    assert "| 设备 | 状态 |" in blocks[0].page_content
